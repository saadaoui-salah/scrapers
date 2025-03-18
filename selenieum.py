import json
import requests
from docx import Document
from docx.shared import Inches

# Load JSON file
with open("linkedin.json", "r", encoding="utf-8") as file:
    data = json.load(file)

# Create a new Word document
doc = Document()

# Iterate through JSON data
for i, post in enumerate(data, start=1):
    title = f"Post {i}"
    text = post.get("text", "")
    image_url = post.get("image", "")

    # Add title
    doc.add_heading(title, level=1)

    # Add text
    doc.add_paragraph(text)

    # Download and insert image
    
# Save the document
doc.save("output.docx")

print("Word document created successfully!")